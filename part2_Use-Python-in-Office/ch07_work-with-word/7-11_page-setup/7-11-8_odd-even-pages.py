from docx import Document
from docx.shared import Cm
doc = Document()

doc.settings.odd_and_even_pages_header_footer = True

section = doc.sections[0]

section.header.paragraphs[0].text = "odd pages' header"
section.footer.paragraphs[0].text = "odd pages' footer"

even_page_header = section.even_page_header
even_page_footer = section.even_page_footer

print(type(even_page_header))
print(type(even_page_footer))

even_page_header.paragraphs[0].text = "even pages' header"
even_page_footer.paragraphs[0].text = "even pages' footer"

for i in range (1,10):
    doc.add_page_break()
    
for section in doc.sections:
    section.page_height = Cm(7)

doc.save("./7-11-8.docx")