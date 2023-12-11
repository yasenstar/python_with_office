from docx import Document
from docx.shared import Cm

doc = Document()

section = doc.sections[0]

header = section.header
header.paragraphs[0].add_run("This is Header")

footer = section.footer
footer.paragraphs[0].add_run("This is Footer")

print(section.header_distance.cm)
print(section.footer_distance.cm)

section.header_distance = Cm(5)
section.footer_distance = Cm(5)

print(section.header_distance.cm)
print(section.footer_distance.cm)

doc.save("./7-11-7_2.docx")