from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

doc = Document()

doc.add_paragraph("This is section 1's paragraph 1")
doc.add_paragraph("This is section 1's paragraph 2")

s1 = doc.sections[0]
s1.header.paragraphs[0].add_run("This is section 1's header")

footer = s1.footer
footer.paragraphs[0].text = "this is the footer in the middle"
footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

s2 = doc.add_section(start_type = WD_SECTION_START.NEW_PAGE)

s2.header.is_linked_to_previous = False
s2.header.paragraphs[0].add_run("This is Section 2's header")

doc.add_paragraph("This is section 2's paragraph 1")
doc.add_paragraph("This is section 2's paragraph 2")

s2.footer.is_linked_to_previous = True
for section in doc.sections:
    section.page_height = Cm(7)

doc.save("./7-11-7_3.docx")