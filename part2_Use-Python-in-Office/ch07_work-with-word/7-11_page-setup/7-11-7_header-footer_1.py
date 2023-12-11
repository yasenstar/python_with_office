from docx import Document
doc = Document()

section = doc.sections[0]

header = section.header
header.paragraphs[0].add_run("This is Header")

footer = section.footer
footer.paragraphs[0].add_run("This is Footer")

doc.save("./7-11-7_1.docx")