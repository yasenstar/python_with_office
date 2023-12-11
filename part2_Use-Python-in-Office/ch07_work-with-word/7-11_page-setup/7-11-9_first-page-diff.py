from docx import Document
from docx.shared import Cm
doc = Document()

section = doc.sections[0]


section.different_first_page_header_footer = True

section.first_page_header.paragraphs[0].text = "first page header"
section.first_page_footer.paragraphs[0].text = "first page footer"

section.header.paragraphs[0].text = "normal hader"
section.footer.paragraphs[0].text = "normal footer"

doc.add_page_break()
doc.add_page_break()

for section in doc.sections:
    section.page_height = Cm(7)

doc.save("./7-11-9.docx")