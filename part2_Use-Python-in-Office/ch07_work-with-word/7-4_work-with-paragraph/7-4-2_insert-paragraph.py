from docx import Document
doc = Document()
paragraph1 = doc.add_paragraph("this is 1st paragraph")
doc.add_page_break()
paragraph3 = doc.add_paragraph("this is 3rd paragraph")
print(paragraph1.text, paragraph3.text)
paragraph2 = paragraph3.insert_paragraph_before("now this is 2nd")
for p in doc.paragraphs:
    print(p.text)
print(len(doc.paragraphs))
doc.save("./7-4-2.docx")