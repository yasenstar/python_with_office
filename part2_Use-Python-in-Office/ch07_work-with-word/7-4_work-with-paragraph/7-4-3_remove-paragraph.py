from docx import Document
doc = Document("./7-4-2.docx")

print(len(doc.paragraphs))

for paragraph in doc.paragraphs:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
print(len(doc.paragraphs))

doc.save("./7-4-3.docx")