from docx import Document
from openpyxl import Workbook

# Word
print("===Word part===")
doc = Document()
print(doc.paragraphs)
print(type(doc.paragraphs))
paragraph1 = doc.add_paragraph("first paragraph")
print(doc.paragraphs)
print(type(paragraph1))
paragraph2 = doc.add_paragraph("this is the content")
print(doc.paragraphs)
print(paragraph2.text)
paragraph2.text = "here is the changed text \n 2nd line of this paragrage"
print(paragraph2.text)
paragraph2.clear()
print(doc.paragraphs)
doc.save("./test.docx")

# Excel
print("===Excel part===")
wb = Workbook()
print(wb.worksheets)
print(type(wb.worksheets))
ws = wb.active
sheet2 = wb.create_sheet("Shee2")
print(wb.worksheets)
print(type(sheet2))
ws.append(["ABC", "BCD", "EFT"])
ws.append([123, "month", 3434])
row_cells = ws[1]
col_cells = ws["B"]
for cell in row_cells:
    print(cell.value)
wb.save("./test.xlsx")