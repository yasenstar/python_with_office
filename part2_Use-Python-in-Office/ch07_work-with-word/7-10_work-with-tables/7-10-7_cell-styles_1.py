from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt

doc = Document()

table = doc.add_table(3,2)
cell = table.cell(1,0)

paragraph = cell.paragraphs[0]
paragraph.text = "Office"
run = paragraph.add_run("遇上了")

run.font.color.rgb = RGBColor.from_string("FF0000")

paragraph.add_run("Python")

paragraph.style.font.size = Pt(15)

paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save("./7-10-7_1.docx")