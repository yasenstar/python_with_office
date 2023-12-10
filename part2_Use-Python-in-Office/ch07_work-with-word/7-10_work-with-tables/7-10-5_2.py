from docx import Document
doc = Document()

table = doc.add_table(3,2)

cell = table.cell(1,0)

paragraph = cell.paragraphs[0]

paragraph.add_run().add_picture("./stamp_bigdata-asso.png")

doc.save("./7-10-5_2.docx")