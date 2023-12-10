from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

doc = Document()

table = doc.add_table(3,2)

table.rows[1].height = Cm(3)
cell = table.cell(1,0)

cell.text = "Office遇上了Python"

paragraph = cell.paragraphs[0]
paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save("./7-10-7_2.docx")