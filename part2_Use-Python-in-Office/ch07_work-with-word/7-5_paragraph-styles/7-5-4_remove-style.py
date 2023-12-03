from docx import Document
doc = Document()

p1 = doc.add_paragraph("a list paragraph", style="List")
p2 = doc.add_paragraph("a list 2nd line", "List 2")

print(type(p2.style))
print(p2.style.name)

p2.style.delete()
print(p2.style.name)

doc.save("./7-5-4.docx")