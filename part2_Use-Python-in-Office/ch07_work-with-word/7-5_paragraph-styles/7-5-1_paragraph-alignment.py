from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as ALIGNMENT

doc = Document()

paragraph1 = doc.add_paragraph("this is left alignment")
paragraph1.alignment = ALIGNMENT.LEFT

paragraph2 = doc.add_paragraph("this is middle alignment")
paragraph2.alignment = ALIGNMENT.CENTER

paragraph3 = doc.add_paragraph("this is right alignment")
paragraph3.alignment = ALIGNMENT.RIGHT

paragraph4 = doc.add_paragraph("this is justified alignment")
paragraph4.alignment = ALIGNMENT.JUSTIFY

paragraph5 = doc.add_paragraph("this is distributed alignment")
paragraph5.alignment = ALIGNMENT.DISTRIBUTE

paragraph6 = doc.add_paragraph("this is Justitied Med alignment")
paragraph6.alignment = ALIGNMENT.JUSTIFY_MED

paragraph7 = doc.add_paragraph("this is Justitied LOW alignment")
paragraph7.alignment = ALIGNMENT.JUSTIFY_LOW

paragraph8 = doc.add_paragraph("this is Justitied HI alignment")
paragraph8.alignment = ALIGNMENT.JUSTIFY_HI

paragraph9 = doc.add_paragraph("this is Thai Justified alignment")
paragraph9.alignment = ALIGNMENT.THAI_JUSTIFY


doc.save("./7-5-1.docx")