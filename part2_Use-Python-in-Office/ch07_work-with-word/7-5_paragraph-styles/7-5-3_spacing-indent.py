from docx import Document
from docx.shared import Pt

doc = Document()

paragraph1 = doc.add_paragraph("this is one line of words")

paragraph1.paragraph_format.space_before = Pt(30)

paragraph1.paragraph_format.space_after = Pt(50)

paragraph2 = doc.add_paragraph("this is 2nd line of words")

# paragraph2.paragraph_format.line_spacing = Pt(50)

paragraph2.paragraph_format.line_spacing = 6


paragraph3 = doc.add_paragraph("this is 3rd line of words")

paragraph2.paragraph_format.first_line_indent = Pt(20)

doc.save("./7-5-3.docx")