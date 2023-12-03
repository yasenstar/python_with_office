from docx import Document
doc = Document()

# normal list
paragraph1 = doc.add_paragraph("有序段落List")
print(paragraph1.style.name)

paragraph1.style = "List"
print(paragraph1.style.name)

paragraph2 = doc.add_paragraph("有序段落List 2", style="List")
print(paragraph2.style.name)
paragraph3 = doc.add_paragraph("有序段落List 3 ", style="List")
print(paragraph3.style.name)

# numbered list

doc.add_paragraph("有序段落List Number 1", style="List Number")
doc.add_paragraph("有序段落List Number 2", style="List Number 2")
doc.add_paragraph("有序段落List Number 2", style="List Number")

# bullet list

doc.add_paragraph("有序段落List Bullet 1", style="List Bullet")
doc.add_paragraph("有序段落List Bullet 2", style="List Bullet 2")
doc.add_paragraph("有序段落List Bullet 3", style="List Bullet 3")

doc.save("./7-5-2.docx")