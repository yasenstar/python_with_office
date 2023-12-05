from docx import Document

doc = Document()

doc.add_heading("certificate")
doc.add_paragraph("this is the sentense")

doc.save("./7-8-2.docx")

# definition for add_heading:

# def add_heading(self,text="",level=1):
#     if not 0 <= level <= 9:
#          raise ValueError("level must be in range 0-9,got %d" %
#  level)
#     style = "Title" if level == 0 else "Heading %d" % level
#     return self.add_paragraph(text,style)