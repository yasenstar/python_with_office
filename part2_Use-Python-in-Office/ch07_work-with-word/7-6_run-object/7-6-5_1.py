from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

# define variables
college = "化学化工"
major = "化学"
class_ = "2201"
name = "冰冰"
join_year = "2021-2022"
association = "大数据协会"
prove_date = "2022年12月30日"

doc = Document()

run_list = list()

# title = doc.add_paragraph("证明")
title = doc.add_heading("证明",0)

title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# print(len(title.runs))
title_run = title.runs[0]
title_run.font.size = Pt(22)
title_run.font.name = "Times New Roman"
title_run.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

body = doc.add_paragraph()
body.add_run("兹证明，")
run_list.append(body.add_run(f"  {college}  "))
body.add_run("学院")
run_list.append(body.add_run(f"  {major}  "))
body.add_run("系")
run_list.append(body.add_run(f" {class_} "))
body.add_run("班")
run_list.append(body.add_run(f"  {name}  "))
body.add_run(f"同学于{join_year}学年加入")
run_list.append(body.add_run(f"{association}"))
body.add_run(
    "（社团名称），鉴于该同学在我社期间各方面表现优秀，特此证明，"
    "并建议按照综测细则给予该同学相应综合量化成绩加分。"
    )

for run in run_list:
    run.font.underline = True

font_size = 14
body.style.font.size = Pt(font_size)
body.paragraph_format.first_line_indent = Pt(font_size * 2)
body.paragraph_format.line_spacing = 2
body.style.font.name = "Times New Roman"
body.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

p_association = doc.add_paragraph(f"{association}")
p_association.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p_association.style.font.name = "Times New Roman"
p_association.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

p_prove_date = doc.add_paragraph(f"{prove_date}")
p_prove_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p_prove_date.style.font.name = "Times New Roman"
p_prove_date.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

doc.save("./7-6-5_1.docx")