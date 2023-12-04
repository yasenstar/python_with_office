from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

doc = Document()

paragraph1 = doc.add_paragraph("Office")

run = paragraph1.add_run("遇上了Python")

# print(type(run.font))

run.font.name = "Aharoni"
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

run.font.size = Pt(50)
run.font.bold = True
run.font.italic = True
# run.font.underline = True
run.font.underline = WD_UNDERLINE.DOUBLE

doc.save("./7-6-4.docx")