from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

# Read Student List from Prepared Excel file
def read_excel(path):
    wb = load_workbook(path)
    ws = wb.active
    student_list = list()
    for i, row in enumerate(ws.rows):
        if i == 0:
            continue
        stu_dict = {
            "name": row[0].value,
            "college": row[1].value,
            "major": row[2].value,
            "class": row[3].value,
            "association": row[4].value,
            "join_year": row[5].value,
            "prove_date": row[6].value,
        }
        if isinstance(row[6].value, int):
            raise Exception("请先在Excel把日期转换成文本格式")
        student_list.append(stu_dict)
    return student_list

def save_template(stu_dict):
    if not stu_dict or not isinstance(stu_dict, dict):
        print("argument error, please give one Dictionary type")
        return
    
    doc = Document()

    run_list = list()

    title = doc.add_paragraph("证明")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # print(len(title.runs))
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.name = "Times New Roman"
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

    body = doc.add_paragraph()
    body.add_run("兹证明，")
    run_list.append(body.add_run(f"  {stu_dict['college']}  "))
    body.add_run("学院")
    run_list.append(body.add_run(f"  {stu_dict['major']}  "))
    body.add_run("系")
    run_list.append(body.add_run(f" {stu_dict['class']} "))
    body.add_run("班")
    run_list.append(body.add_run(f"  {stu_dict['name']}  "))
    body.add_run(f"同学于{stu_dict['join_year']}学年加入")
    run_list.append(body.add_run(f"{stu_dict['association']}"))
    body.add_run(
        "（社团名称），鉴于该同学在我社期间各方面表现优秀，特此证明，"
        "并建议按照综测细则给予该同学相应综合量化成绩加分。"
        )

    for run in run_list:
        run.font.underline = True

    font_size = 14
    body.style.font.size = Pt(font_size)
    body.paragraph_format.first_line_indent = Pt(font_size * 2)
    body.paragraph_format.line_spacing = 2
    body.style.font.name = "Times New Roman"
    body.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

    p_association = doc.add_paragraph(f"{stu_dict['association']}")
    p_association.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_association.style.font.name = "Times New Roman"
    p_association.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

    p_prove_date = doc.add_paragraph(f"{stu_dict['prove_date']}")
    p_prove_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_prove_date.style.font.name = "Times New Roman"
    p_prove_date.style.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')

    doc.save(f"./output/7-6-5_2_{stu_dict['name']}_参加社团证明.docx")

if __name__ == '__main__':
    students = read_excel("./7-6-5.xlsx")
    print(students)
    for student in students:
        save_template(student)