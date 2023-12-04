from docx import Document
doc = Document()

paragraph = doc.add_paragraph("Office")

run1 = paragraph.add_run("遇上了")

print(run1.text)

run1.clear() # same as doc.paragraphs[0].runs[1].clear()

print(run1.text)

doc.save("./7-6-3.docx")