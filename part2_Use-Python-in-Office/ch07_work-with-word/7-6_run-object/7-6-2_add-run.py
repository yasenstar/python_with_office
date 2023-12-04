from docx import Document
doc = Document()

paragraph = doc.add_paragraph("Office")

# print(len(paragraph.runs))
# print(paragraph.runs[0].text)

run1 = paragraph.add_run("遇上了")
run2 = paragraph.add_run("Python")

print(len(paragraph.runs))

print(run1.text)

print(doc.paragraphs[0].runs[2].text)

doc.save("./7-6-2.docx")