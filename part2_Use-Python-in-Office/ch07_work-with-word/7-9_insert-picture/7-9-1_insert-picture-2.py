from docx import Document
img_path = "./galaxy.jpg"
doc = Document()

# same paragraph, same line
doc.add_picture(img_path)
doc.paragraphs[0].runs[0].add_picture(img_path)

# same paragraph, different lines
p = doc.add_paragraph()
p.add_run().add_picture(img_path)
run = p.add_run("\n")
p.add_run().add_picture(img_path)

# different paragraph, different lines
doc.add_paragraph().add_run().add_picture(img_path)
doc.add_paragraph().add_run().add_picture(img_path)

doc.save("./7-9-1_insert-2.docx")