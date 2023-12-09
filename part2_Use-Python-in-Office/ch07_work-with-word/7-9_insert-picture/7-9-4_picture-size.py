from docx import Document
from docx.shared import Mm, Cm, Pt

img_path1 = "./galaxy.jpg"
img_path2 = "./stamp_bigdata-asso.png"
doc = Document()

img1 = doc.add_picture(img_path1, width=Mm(50), height=Mm(50))

doc.add_picture(img_path2, width=Cm(5), height=Cm(5))

doc.add_picture(img_path1, width=Cm(5))

doc.add_picture(img_path2, height=Pt(141.73))

print(img1.width.cm)
print(img1.width.pt)

img1.width = Cm(8)
print(img1.width.cm)

doc.save("./7-9-4.docx")