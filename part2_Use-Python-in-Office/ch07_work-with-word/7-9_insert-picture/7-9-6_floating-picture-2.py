from docx import Document
from docx.oxml import parse_xml, CT_Picture, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.shared import Cm

class CT_Anchor(BaseOxmlElement):
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')
    @classmethod
    def new(cls, pos_x, pos_y, cx, cy, pic_id, pic, behind):
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y, behind))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = pic_id
        anchor.docPr.name = 'Picture %d'% pic_id
        anchor.graphic.graphicData.uri = ('http://schemas.openxmlformats.org/drawingml/2006/picture')
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor
    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y, behind):
        pic_id = 0
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(pos_x, pos_y, cx, cy, shape_id, pic, behind)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor
    @classmethod
    def _anchor_xml(cls, pos_x, pos_y, behind):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" '
            ' simplePos="0" relativeHeight="0 " \n'
            ' behindDoc="{}" locked="0" layoutInCell="1" '
            ' allowOverlap="1" \n {}>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>{}</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>{}</wp:posOffset>\n'
            '  </wp:positionV>\n'
            '  <wp:extent cx="2695575" cy="1619250"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="26429" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>'.format(behind, nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y)
            )
        )
        
def add_anchor_picture(p, img, width=None, height=None, pos_x=0, pos_y=0, behind=0):
    run = p.add_run()
    r_id, image = run.part.get_or_add_image(img)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, name = run.part.next_id, image.filename
    anchor = CT_Anchor.new_pic_anchor(shape_id, r_id, name, cx, cy, pos_x, pos_y, behind)
    run._r.add_drawing(anchor)

register_element_cls('wp:anchor', CT_Anchor)

doc = Document("./7-6-5_2_冰冰_参加社团证明.docx")

paragraph = doc.add_paragraph()
img_path = "./stamp_bigdata-asso.png"
# doc.add_picture(img_path, Cm(5), Cm(5))

add_anchor_picture(paragraph, img_path, width=Cm(3), pos_x=Cm(15.5), pos_y=Cm(8))

doc.save("./7-9-6.docx")