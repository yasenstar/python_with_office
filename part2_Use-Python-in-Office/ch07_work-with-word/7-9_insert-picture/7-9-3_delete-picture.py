from docx import Document
img_path = "./galaxy.jpg"
doc = Document()

paragraph = doc.add_paragraph()
for i in range(1,10):
    paragraph.add_run().add_picture(img_path)

run = paragraph.add_run()
run.add_picture(img_path)

print(len(doc.inline_shapes._inline_lst))

run.clear()

print(len(doc.inline_shapes._inline_lst))

paragraph.clear()

print(len(doc.inline_shapes._inline_lst))