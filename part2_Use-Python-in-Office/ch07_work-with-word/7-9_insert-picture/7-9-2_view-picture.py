from docx import Document
img_path = "./galaxy.jpg"
doc = Document("./7-9-1_insert-2.docx")

doc.add_picture(img_path)
doc.add_picture(img_path)
doc.add_picture(img_path)

for shape in doc.inline_shapes:
    print(shape, shape.type)

shape_list = doc.inline_shapes._inline_lst
print(len(shape_list))

# doc.save("./7-9-2_1.docx")