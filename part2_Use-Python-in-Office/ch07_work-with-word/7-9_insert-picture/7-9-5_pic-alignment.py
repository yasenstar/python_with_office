from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm

img_path = "./stamp_bigdata-asso.png"

doc = Document()

doc.add_picture(img_path,width=Mm(50),height=Mm(50))

doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save("./7-9-5.docx")